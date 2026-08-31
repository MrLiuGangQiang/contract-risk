"""合同风险识别业务用例（《11-合同风险识别核心功能设计》）。

上传合同 → 文本解析 → 当前用户生效规则关键词匹配 → 风险结果落库。
同步执行；文件大小/类型校验在 API 与 Service 双层实施。
"""
import asyncio
import base64
import io
import json
import logging
import re
import time
import uuid

import httpx
import pymupdf
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Awaitable, Callable

from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.exceptions import BizException
from app.domain.constants import (
    CONTRACT_ALLOWED_EXTENSIONS,
    CONTRACT_MAX_SIZE,
    CONTRACT_STATUS_DONE,
    CONTRACT_STATUS_FAILED,
    CONTRACT_STATUS_SCANNING,
)

RISK_RULE_SEVERITIES = {"high", "medium", "low"}
from app.models.contract import Contract, ContractRisk
from app.repositories.contract import ContractRepository
from app.repositories.logs import OperationLogRepository
from app.schemas.contract import ContractDetailOut, ContractOut, ContractPageOut, ContractRiskOut
from app.services.ai_config_service import AIConfigService
from app.services.risk_rule_service import RiskRuleService

logger = logging.getLogger(__name__)

_SEVERITY_ORDER = {"high": 0, "medium": 1, "low": 2}
_SNIPPET_RADIUS = 60

# 扫描版 PDF 识别参数（《11》第 4.1 节）
SCANNED_PDF_TEXT_THRESHOLD = 50  # 文本层少于该字符数视为扫描版（纯图片）
_VISION_PAGE_CONCURRENCY = 4  # 视觉转录并发页数
_VISION_PAGE_DPI = 140  # 页面渲染 DPI（清晰度与体积平衡）
_VISION_PAGE_TIMEOUT_SECONDS = 120  # 单页转录超时
# AI 逐条规则校验并发数（每条规则 = 一个独立 AI 任务）
_RULE_CHECK_CONCURRENCY = 4

_RULE_CHECK_SYSTEM = (
    "你是代表乙方（供应商/服务方）的资深合同法律顾问。乙方是提供产品或服务、向甲方交付并收款的一方，"
    "甲方是采购方、付款方。请始终站在乙方立场，理解下面这条风险规则想保护乙方什么，"
    "再逐段检查合同文本是否对乙方不利、加重乙方义务、扩大乙方责任、克扣乙方回款或侵占乙方权利。"
    "只返回一个 JSON 对象，不要输出其它内容。"
)

_VISION_OCR_PROMPT = (
    "请把这张文档页面图片中的文字完整转录出来，保持原有段落与编号，"
    "只输出转录的文字内容，不要添加任何解释或评论。"
)


def _needs_vision_ocr(ext: str, text: str) -> bool:
    """是否需要视觉 OCR（仅扫描版 PDF：文本层几乎为空）。"""
    return ext == "pdf" and len(text.strip()) < SCANNED_PDF_TEXT_THRESHOLD


def _decode_text(data: bytes) -> str:
    """Decode txt with UTF-8 then GB18030 fallback."""
    for encoding in ("utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _parse_document(ext: str, data: bytes) -> str:
    """Extract plain text from txt/pdf/docx."""
    if ext == "txt":
        return _decode_text(data)
    if ext == "pdf":
        doc = pymupdf.open(stream=data, filetype="pdf")
        try:
            return "\n".join(page.get_text() for page in doc)
        finally:
            doc.close()
    if ext == "docx":
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    raise BizException(20002, "不支持的文件格式，仅支持 txt/pdf/docx")


def _risk_out(risk: ContractRisk) -> ContractRiskOut:
    """ORM -> DTO."""
    return ContractRiskOut(
        id=risk.id,
        rule_code=risk.rule_code,
        snippet_start=risk.snippet_start,
        snippet_end=risk.snippet_end,
        rule_name=risk.rule_name,
        category=risk.category,
        severity=risk.severity,
        matched_keywords=risk.matched_keywords or [],
        snippet=risk.snippet,
        description=risk.description,
        suggestion=risk.suggestion,
        sort_order=risk.sort_order,
        risk_source=risk.risk_source or "rule",
        created_at=risk.created_at,
    )


def _contract_out(contract: Contract) -> ContractOut:
    """ORM -> DTO."""
    return ContractOut(
        id=contract.id,
        user_id=contract.user_id,
        file_name=contract.file_name,
        file_ext=contract.file_ext,
        file_size=contract.file_size,
        total_chars=contract.total_chars,
        status=contract.status,
        risk_count=contract.risk_count,
        high_count=contract.high_count,
        medium_count=contract.medium_count,
        low_count=contract.low_count,
        created_at=contract.created_at,
    )


_AI_RESULT_MARKERS = ("== 结果 ==", "==结果==")


def split_ai_output(content: str) -> tuple[str, str]:
    """拆分 AI 两段式输出：(分析思路, 结果 JSON 原文)。

    未出现结果分隔标记时，整体视为分析思路（结构化结果为空）。
    """
    for marker in _AI_RESULT_MARKERS:
        index = content.find(marker)
        if index != -1:
            return content[:index].strip(), content[index + len(marker):]
    return content.strip(), ""


def _parse_ai_json(content: str) -> list[dict[str, Any]]:
    """Extract a JSON array from an AI response (best effort)."""
    _, result_part = split_ai_output(content)
    source = result_part if result_part.strip() else content
    start = source.find("[")
    end = source.rfind("]")
    if start == -1 or end == -1 or end <= start:
        return []
    try:
        data = json.loads(source[start : end + 1])
    except ValueError:
        return []
    if not isinstance(data, list):
        return []
    findings: list[dict[str, Any]] = []
    for i, item in enumerate(data[:100]):
        if not isinstance(item, dict):
            continue
        category = str(item.get("category") or "").strip()[:32] or "general"
        severity = str(item.get("severity") or "medium")
        if severity not in RISK_RULE_SEVERITIES:
            severity = "medium"
        code = str(item.get("rule_code") or "").strip() or f"AI_FINDING_{i + 1}"
        findings.append(
            {
                "rule_code": code[:64],
                "rule_name": str(item.get("rule_name") or f"AI Finding {i + 1}")[:128],
                "category": category,
                "severity": severity,
                "matched_keywords": [str(k) for k in (item.get("matched_keywords") or []) if str(k)][:20],
                "snippet": str(item.get("snippet") or "")[: _SNIPPET_RADIUS * 4],
                "description": str(item.get("description") or "")[:2000],
                "suggestion": str(item.get("suggestion") or "")[:2000],
                "sort_order": 10000 + i,
                "risk_source": "ai",
            }
        )
    return findings


def _split_contract_chunks(text: str, max_chars: int = 1200) -> list[dict[str, Any]]:
    """把合同按条款边界切块，返回 [{text, start, end}]（原文保真 + 字符偏移）。

    优先级：空行 → 条款编号（第X条/X.X/一、二、）→ 句号 → 硬切。
    """
    # 先按空行粗分（记录原始偏移）
    blocks: list[dict[str, Any]] = []
    pos = 0
    for m in re.finditer(r"\n\s*\n", text):
        block = text[pos:m.start()]
        if block.strip():
            blocks.append({"text": block, "start": pos, "end": m.start()})
        pos = m.end()
    if text[pos:].strip():
        blocks.append({"text": text[pos:], "start": pos, "end": len(text)})
    if not blocks:
        blocks = [{"text": text, "start": 0, "end": len(text)}]

    # 再按条款编号细分
    clause_pattern = re.compile(
        r"(?=\n\s*(?:第[一二三四五六七八九十百0-9]+[条节章款]|[0-9]{1,2}\.[0-9]{1,2}|[一二三四五六七八九十]+[、.]))"
    )
    chunks: list[dict[str, Any]] = []
    for block in blocks:
        parts = [p for p in clause_pattern.split(block["text"]) if p.strip()]
        if len(parts) <= 1 and len(block["text"]) > max_chars:
            parts = [p for p in re.split(r"(?<=[。；;])", block["text"]) if p.strip()]
        offset = block["start"]
        # 重建每部分在原文中的偏移
        search_from = block["start"]
        for part in parts:
            idx = text.find(part.strip(), search_from)
            if idx == -1:
                idx = search_from
            p_start = idx
            p_text = part.strip()
            while len(p_text) > max_chars:
                chunks.append({"text": p_text[:max_chars], "start": p_start, "end": p_start + max_chars})
                p_text = p_text[max_chars:]
                p_start += max_chars
            if p_text:
                chunks.append({"text": p_text, "start": p_start, "end": p_start + len(p_text)})
            search_from = p_start + len(p_text)
    return chunks or [{"text": text, "start": 0, "end": len(text)}]


def _parse_clause_list(content: str) -> list[dict[str, str]]:
    """解析条款摘要 JSON 数组（宽容处理）。"""
    start = content.find("[")
    end = content.rfind("]")
    if start == -1 or end == -1 or end <= start:
        return []
    try:
        data = json.loads(content[start : end + 1])
    except ValueError:
        return []
    if not isinstance(data, list):
        return []
    result: list[dict[str, str]] = []
    for item in data:
        if isinstance(item, dict):
            result.append({
                "title": str(item.get("title") or "条款")[:60],
                "summary": str(item.get("summary") or "")[:200],
            })
    return result


def _locate_by_rule_keywords(text: str, rule_text: str) -> int:
    """用规则中的实义中文词在原文中定位（兜底：缺失条款类风险定位到相关区域）。"""
    stopwords = {
        "必须", "应当", "不得", "应该", "需要", "要", "明确", "约定", "完整", "合理",
        "符合", "提供", "保证", "确保", "避免", "防止", "禁止", "的", "和", "或", "与",
    }
    # 取规则中较长的连续词（2~8 字）作为定位词，过滤停用词与虚词
    candidates = re.findall(r"[\u4e00-\u9fa5]{2,8}", rule_text)
    best: tuple[int, int] = (-1, 0)
    for word in candidates:
        if word in stopwords:
            continue
        if len(word) < 2:
            continue
        pos = text.find(word)
        if pos != -1 and len(word) > best[1]:
            best = (pos, len(word))
    return best[0]


def _parse_single_rule_result(content: str) -> dict[str, Any]:
    """解析单条规则校验的 JSON 结果（宽容处理）。"""
    start = content.find("{")
    end = content.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return {}
    try:
        data = json.loads(content[start : end + 1])
    except ValueError:
        return {}
    return data if isinstance(data, dict) else {}



class ContractRiskService:

    """合同风险识别业务用例。"""

    def __init__(self, session: AsyncSession) -> None:
        self._session = session
        self._repo = ContractRepository(session)
        self._risk_rule_service = RiskRuleService(session)
        self._ai_config_service = AIConfigService(session)
        self._audit_repo = OperationLogRepository(session)

    async def register_upload(
        self,
        *,
        user_id: int,
        file_name: str,
        content: bytes,
        request_meta: dict[str, Any],
    ) -> Contract:
        """上传登记：校验 → 保存文件 → 解析文本 → 合同以「扫描中」状态入库并提交。

        扫描在后台任务中执行（支持关闭页面继续跑），解析失败不入库（《11》第 4 节）。
        """
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if ext not in CONTRACT_ALLOWED_EXTENSIONS:
            raise BizException(20002, "仅支持 txt/pdf/docx 文件")
        if not content:
            raise BizException(20002, "文件内容为空")
        if len(content) > CONTRACT_MAX_SIZE:
            raise BizException(20002, "文件不能超过 20MB")

        text = await asyncio.to_thread(_parse_document, ext, content)

        settings = get_settings()
        upload_path = Path(settings.upload_dir)
        upload_path.mkdir(parents=True, exist_ok=True)
        stored_name = f"{uuid.uuid4().hex}.{ext}"
        stored_path = upload_path / stored_name
        stored_path.write_bytes(content)

        contract = Contract(
            user_id=user_id,
            file_name=file_name[:255],
            file_path=str(stored_path),
            file_ext=ext,
            file_size=len(content),
            text_content=text,
            total_chars=len(text),
            status=CONTRACT_STATUS_SCANNING,
            created_by=user_id,
            updated_by=user_id,
        )
        await self._repo.add_contract(contract)
        await self._audit(
            "contract.upload", "POST", "/api/v1/contracts/upload",
            {"file_name": file_name}, user_id, request_meta,
        )
        await self._session.commit()
        logger.info("contract registered for background scan", extra={"user_id": user_id, "contract_id": contract.id})
        return contract

    async def run_scan(
        self,
        *,
        user_id: int,
        contract_id: int,
        request_meta: dict[str, Any],
        progress_cb: Callable[[int, str], Awaitable[None]] | None = None,
        init_cb: Callable[[dict[str, int]], Awaitable[None]] | None = None,
        task_cb: Callable[[str, str, int], Awaitable[None]] | None = None,
        ai_chunk_cb: Callable[[str], Awaitable[None]] | None = None,
        ai_status_cb: Callable[[str, int], Awaitable[None]] | None = None,
        rule_status_cb: Callable[[str, str, str], Awaitable[None]] | None = None,
        action: str = "contract.scan",
        audit_path: str = "/api/v1/contracts/scan",
    ) -> ContractDetailOut:
        """执行扫描（《11》2.2）：条款拆分 + 逐条规则并发校验 → 生成报告。

        AI 未配置时抛业务异常（扫描依赖 AI 按内容理解规则）。
        """

        async def notify(progress: int, stage: str) -> None:
            if progress_cb is not None:
                await progress_cb(progress, stage)

        contract = await self._get_or_404(user_id, contract_id)
        await notify(10, "正在读取合同内容")
        text = contract.text_content

        # 扫描版 PDF（无文本层）：先执行 AI 视觉转录，再进入风险识别
        if _needs_vision_ocr(contract.file_ext, text):
            await notify(12, "检测到扫描版 PDF（无文本层），启动 AI 视觉识别")

            async def vision_progress(idx: int, total: int, chars: int) -> None:
                percent = 12 + int(40 * idx / max(total, 1))
                await notify(percent, f"AI 视觉识别第 {idx}/{total} 页（累计转录 {chars} 字符）")

            data = Path(contract.file_path).read_bytes()
            text = await self._vision_ocr_pdf(data, page_cb=vision_progress)
            contract.text_content = text
            contract.total_chars = len(text)
            await self._session.flush()
            await notify(52, f"视觉识别完成：共转录 {len(text)} 字符")

        rules = await self._risk_rule_service.list_effective(user_id)
        enabled_rules = [r for r in rules if r.enabled]

        # 维度任务清单（纯展示分组；实际校验是逐条规则并发）
        rules_by_category: dict[str | None, list[Any]] = {}
        for rule in enabled_rules:
            rules_by_category.setdefault(rule.category, []).append(rule)
        if init_cb is not None:
            await init_cb({cat or "未分类": len(rs) for cat, rs in rules_by_category.items()})
        await notify(18, f"已准备 {len(enabled_rules)} 条规则，开始理解合同结构并逐条并发校验")

        # AI 两阶段校验：条款拆分 → 逐条规则并发判定
        if ai_status_cb is not None:
            await ai_status_cb("running", 0)
        ai_hits, ai_status = await self._analyze_rules_with_ai(
            text,
            enabled_rules,
            chunk_cb=ai_chunk_cb,
            rule_status_cb=rule_status_cb,
        )
        if ai_status_cb is not None:
            await ai_status_cb(ai_status, len(ai_hits))
        if ai_status == "skipped":
            raise BizException(20002, "AI 大模型未配置，请先在「系统配置 → AI 大模型」启用后重试")

        # 维度命中统计（报告按维度筛选展示）
        dim_hits: dict[str | None, int] = {}
        for hit in ai_hits:
            dim_hits[hit["category"]] = dim_hits.get(hit["category"], 0) + 1
        if task_cb is not None:
            for cat in rules_by_category:
                await task_cb(cat or "未分类", "done", dim_hits.get(cat, 0))

        await notify(90, "正在汇总生成风险报告")
        risks = await self._save_risks(contract, user_id, ai_hits)
        await self._update_contract_counts(contract, ai_hits)
        contract.status = CONTRACT_STATUS_DONE
        await notify(100, "扫描完成，报告已生成")
        await self._audit(action, "POST", audit_path, {"risk_count": len(ai_hits)}, user_id, request_meta)
        await self._session.commit()
        logger.info("contract scan finished", extra={"user_id": user_id, "contract_id": contract_id})
        return ContractDetailOut(contract=_contract_out(contract), risks=[_risk_out(r) for r in risks])

    async def mark_failed(self, *, user_id: int, contract_id: int, error: str) -> None:
        """扫描失败：合同置为失败状态（独立事务，后台任务调用）。"""
        contract = await self._repo.get_contract(user_id, contract_id)
        if contract is None:
            return
        contract.status = CONTRACT_STATUS_FAILED
        await self._session.commit()
        logger.warning("contract scan marked failed", extra={"contract_id": contract_id, "error": error})

    async def list_contracts(
        self,
        *,
        user_id: int,
        page: int,
        page_size: int,
        keyword: str | None,
        severity: str | None,
    ) -> ContractPageOut:
        """合同分页列表。"""
        items, total = await self._repo.list_contracts(
            user_id=user_id, page=page, page_size=page_size, keyword=keyword, severity=severity
        )
        return ContractPageOut(
            items=[_contract_out(c) for c in items], total=total, page=page, page_size=page_size
        )

    async def get_own_contract(self, user_id: int, contract_id: int) -> Contract | None:
        """按用户查询合同归属（用于后台任务进度查询鉴权）。"""
        return await self._repo.get_contract(user_id, contract_id)

    async def get_detail(self, *, user_id: int, contract_id: int) -> ContractDetailOut:
        """合同详情 + 风险列表。"""
        contract = await self._get_or_404(user_id, contract_id)
        risks = await self._repo.list_risks(contract.id)
        return ContractDetailOut(contract=_contract_out(contract), risks=[_risk_out(r) for r in risks])

    async def register_rescan(
        self,
        *,
        user_id: int,
        contract_id: int,
        request_meta: dict[str, Any],
    ) -> Contract:
        """登记重新扫描：校验后合同置为「扫描中」并提交；扫描在后台任务执行。"""
        contract = await self._get_or_404(user_id, contract_id)
        if contract.status == CONTRACT_STATUS_SCANNING:
            raise BizException(10000, "该合同正在扫描中，请稍后")
        contract.status = CONTRACT_STATUS_SCANNING
        await self._audit(
            "contract.rescan", "POST", f"/api/v1/contracts/{contract_id}/rescan",
            {}, user_id, request_meta,
        )
        await self._session.commit()
        return contract

    async def delete(self, *, user_id: int, contract_id: int, request_meta: dict[str, Any]) -> None:
        """软删除合同（并删除风险结果）。"""
        contract = await self._get_or_404(user_id, contract_id)
        await self._repo.soft_delete_contract(contract)
        await self._audit(
            "contract.delete", "DELETE", f"/api/v1/contracts/{contract_id}",
            {"file_name": contract.file_name}, user_id, request_meta,
        )
        await self._session.commit()
        logger.info("contract deleted", extra={"user_id": user_id, "contract_id": contract_id})

    # ==================== 私有 ====================

    async def _vision_ocr_pdf(
        self,
        data: bytes,
        page_cb: Callable[[int, int, int], Awaitable[None]] | None = None,
    ) -> str:
        """扫描版 PDF 视觉转录：逐页渲染为图片，AI 视觉模型并发转录后拼接。

        未配置/未启用 AI 时抛出业务异常（扫描版无文本层，规则引擎无从工作）。
        """
        cfg = await self._ai_config_service.get_plain()
        enabled = bool(cfg.get("enabled", False))
        api_key = str(cfg.get("api_key_enc", "") or "")
        api_base = str(cfg.get("api_base", "") or "")
        model = str(cfg.get("model", "") or "")
        if not enabled or not api_key or not api_base or not model:
            raise BizException(
                20002,
                "该 PDF 为扫描版（无文本层），需启用 AI 大模型进行视觉识别，请先在系统配置中配置 AI",
            )

        doc = pymupdf.open(stream=data, filetype="pdf")
        try:
            pages_b64 = []
            for page in doc:
                pix = page.get_pixmap(dpi=_VISION_PAGE_DPI)
                pages_b64.append(base64.b64encode(pix.tobytes("png")).decode())
        finally:
            doc.close()
        total_pages = len(pages_b64)
        if total_pages == 0:
            raise BizException(20002, "PDF 无任何页面，无法识别")

        url = f"{api_base.rstrip('/')}/chat/completions"
        sem = asyncio.Semaphore(_VISION_PAGE_CONCURRENCY)
        results: dict[int, str] = {}
        done_pages = {"count": 0, "chars": 0}
        report_lock = asyncio.Lock()

        async def transcribe_page(index: int, image_b64: str) -> None:
            async with sem:
                try:
                    async with httpx.AsyncClient(
                        timeout=_VISION_PAGE_TIMEOUT_SECONDS
                    ) as client:
                        resp = await client.post(
                            url,
                            headers={"Authorization": f"Bearer {api_key}"},
                            json={
                                "model": model,
                                "messages": [
                                    {
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": _VISION_OCR_PROMPT},
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:image/png;base64,{image_b64}"
                                                },
                                            },
                                        ],
                                    }
                                ],
                                "temperature": 0,
                            },
                        )
                        resp.raise_for_status()
                        content = resp.json()["choices"][0]["message"]["content"] or ""
                except Exception as exc:
                    logger.warning("vision OCR page %d failed: %s", index + 1, exc)
                    content = ""
                results[index] = content.strip()
                async with report_lock:
                    done_pages["count"] += 1
                    done_pages["chars"] = sum(len(v) for v in results.values())
                    if page_cb is not None:
                        await page_cb(done_pages["count"], total_pages, done_pages["chars"])

        await asyncio.gather(*(transcribe_page(i, b64) for i, b64 in enumerate(pages_b64)))
        merged = "\n\n".join(results[i] for i in range(total_pages)).strip()
        if not merged:
            raise BizException(20002, "AI 视觉识别未提取到任何文字，请确认文件内容清晰可读")
        return merged

    async def _ai_heartbeat(
        self,
        progress_cb: Callable[[int, str], Awaitable[None]],
        started: float,
    ) -> None:
        """AI 调用期间每 3 秒刷新一次“还在思考”提示，让用户看到进度。"""
        while True:
            await asyncio.sleep(3)
            elapsed = int(time.monotonic() - started)
            await progress_cb(70, f"AI 正在深度分析中（已等待 {elapsed} 秒）...")
    async def _analyze_rules_with_ai(
        self,
        text: str,
        rules: list[Any],
        *,
        chunk_cb: Callable[[str], Awaitable[None]] | None = None,
        rule_status_cb: Callable[[str, str, str], Awaitable[None]] | None = None,
    ) -> tuple[list[dict[str, Any]], str]:
        """AI 两阶段校验（《11》2.2.1）。

        阶段一：合同条款拆分（1 次调用，输出条款摘要清单，本地保留原文 chunk）；
        阶段二：每条规则 × 条款清单并发判定（不重复传全量合同）；
        命中后从原文 chunk 提取证据；AI 未配置返回 skipped（调用方报错）。
        """
        cfg = await self._ai_config_service.get_plain()
        enabled = bool(cfg.get("enabled", False))
        api_key = str(cfg.get("api_key_enc", "") or "")
        api_base = str(cfg.get("api_base", "") or "")
        model = str(cfg.get("model", "") or "")
        if not enabled or not api_key or not api_base or not model:
            return [], "skipped"

        url = f"{api_base.rstrip('/')}/chat/completions"
        timeout = httpx.Timeout(float(cfg.get("timeout_seconds", 60)), connect=10.0)

        # ========== 阶段一：条款拆分 + 摘要 ==========
        chunks = _split_contract_chunks(text)
        summaries = await self._summarize_clauses(
            chunks, url=url, api_key=api_key, model=model, timeout=timeout
        )
        clause_list = "\n".join(
            f"{i + 1}|{s['title']}|{s['summary']}" for i, s in enumerate(summaries)
        )

        # ========== 阶段二：逐条规则并发判定 ==========
        semaphore = asyncio.Semaphore(_RULE_CHECK_CONCURRENCY)

        async def check_one(rule: Any) -> dict[str, Any] | None:
            async with semaphore:
                rule_text = rule.rule_text.strip()
                rid = str(rule.id)
                if rule_status_cb is not None:
                    await rule_status_cb(rid, "running", rule_text)
                user_content = (
                    f"风险规则（乙方视角）：{rule_text}\n\n"
                    f"合同条款清单（编号|标题|摘要）：\n{clause_list}\n\n"
                    "请站在乙方立场判断合同是否违反该规则、是否对乙方不利。注意区分两种情形：\n"
                    "1. 合同有明确条款但内容对乙方不利或违反规则：按对乙方不利程度判定 high/medium/low；\n"
                    "2. 合同完全缺失该规则要求的、对乙方有利的约定（缺失保护条款）：通常判 medium，"
                    "核心致命缺失才判 high。\n"
                    "只返回 JSON："
                    '{"matched":true,"clause_indices":[相关条款编号],"excerpt":"合同原文中的原句",'
                    '"severity":"high|medium|low",'
                    '"explanation":"通俗解释对乙方的风险后果","suggestion":"给乙方的修改建议"}'
                    ' 或 {"matched":false}。命中时 excerpt 必须从合同原文逐字摘录。'
                )
                try:
                    async with httpx.AsyncClient(timeout=timeout) as client:
                        resp = await client.post(
                            url,
                            headers={"Authorization": f"Bearer {api_key}"},
                            json={
                                "model": model,
                                "messages": [
                                    {"role": "system", "content": _RULE_CHECK_SYSTEM},
                                    {"role": "user", "content": user_content},
                                ],
                                "temperature": 0,
                            },
                        )
                        resp.raise_for_status()
                        content = resp.json()["choices"][0]["message"]["content"] or ""
                    payload = _parse_single_rule_result(content)
                    if payload.get("matched") is True:
                        hit = self._build_rule_hit(rule, payload, chunks)
                        if rule_status_cb is not None:
                            await rule_status_cb(rid, "matched", rule_text)
                        return hit
                    if rule_status_cb is not None:
                        await rule_status_cb(rid, "clean", rule_text)
                    return None
                except Exception as exc:
                    logger.warning("AI rule check failed for rule %s: %s", rule.id, exc)
                    if rule_status_cb is not None:
                        await rule_status_cb(rid, "failed", rule_text)
                    return None

        enabled_rules = [r for r in rules if r.enabled]
        results = await asyncio.gather(*(check_one(r) for r in enabled_rules))
        hits = [r for r in results if r is not None]
        if chunk_cb is not None:
            for hit in hits:
                await chunk_cb(
                    f"[{hit['rule_name']}] {hit['severity'].upper()}：{hit['description'][:120]}\n"
                )
        return hits, "done"

    async def _summarize_clauses(
        self,
        chunks: list[dict[str, Any]],
        *,
        url: str,
        api_key: str,
        model: str,
        timeout: httpx.Timeout,
    ) -> list[dict[str, str]]:
        """阶段一：让 AI 为每个条款块生成标题与摘要（失败时降级取块首句）。"""
        numbered = "\n\n".join(f"[{i + 1}] {c['text'][:1500]}" for i, c in enumerate(chunks))
        system = (
            "你是合同结构分析专家。把合同文本拆成条款清单，每个条款给出简短标题和一句话摘要。"
            '只返回 JSON 数组：[{"index":1,"title":"条款标题","summary":"一句话摘要"}]。'
        )
        try:
            async with httpx.AsyncClient(timeout=timeout) as client:
                resp = await client.post(
                    url,
                    headers={"Authorization": f"Bearer {api_key}"},
                    json={
                        "model": model,
                        "messages": [
                            {"role": "system", "content": system},
                            {"role": "user", "content": numbered},
                        ],
                        "temperature": 0,
                    },
                )
                resp.raise_for_status()
                content = resp.json()["choices"][0]["message"]["content"] or ""
            data = _parse_clause_list(content)
            if len(data) == len(chunks):
                return data
        except Exception as exc:
            logger.warning("clause summarization failed, using fallback: %s", exc)
        return [
            {"title": f"第{i + 1}条", "summary": c["text"].strip()[:60]}
            for i, c in enumerate(chunks)
        ]

    def _build_rule_hit(
        self, rule: Any, payload: dict[str, Any], chunks: list[dict[str, Any]]
    ) -> dict[str, Any]:
        """将单条 AI 校验结果映射为风险命中结构（证据含原文偏移，支持局部预览）。"""
        severity = str(payload.get("severity") or "medium")
        if severity not in RISK_RULE_SEVERITIES:
            severity = "medium"
        indices = [int(i) - 1 for i in (payload.get("clause_indices") or []) if str(i).isdigit()]
        indices = [i for i in indices if 0 <= i < len(chunks)]
        if indices:
            snippet = "\n\n".join(chunks[i]["text"][:600] for i in indices)
            start = min(chunks[i]["start"] for i in indices)
            end = max(chunks[i]["end"] for i in indices)
        else:
            # 兜底定位：优先 AI 返回的原文摘录，其次规则文本中的实义词
            full_text = "".join(c["text"] for c in chunks)
            excerpt = str(payload.get("excerpt") or "").strip()
            pos = full_text.find(excerpt) if excerpt else -1
            if pos == -1:
                pos = _locate_by_rule_keywords(full_text, rule.rule_text)
            if pos != -1:
                snippet = full_text[pos:pos + 240]
                start, end = pos, min(pos + 240, len(full_text))
            else:
                snippet, start, end = "", None, None
        return {
            "rule_id": rule.id,
            "rule_name": rule.rule_text.strip()[:128],
            "category": rule.category,
            "severity": severity,
            "matched_keywords": [],
            "snippet": snippet,
            "snippet_start": start,
            "snippet_end": end,
            "description": str(payload.get("explanation") or rule.rule_text.strip())[:2000],
            "suggestion": str(payload.get("suggestion") or "")[:2000],
            "sort_order": rule.sort_order,
            "risk_source": "ai",
        }

    async def _save_risks(self, contract: Contract, user_id: int, hits: list[dict[str, Any]]) -> list[ContractRisk]:

        """写入风险结果并返回实体列表（不提交）。"""
        risks = [
            ContractRisk(
                contract_id=contract.id,
                user_id=user_id,
                rule_id=h.get("rule_id"),
                rule_code=None,
                snippet_start=h.get("snippet_start"),
                snippet_end=h.get("snippet_end"),
                rule_name=h["rule_name"],
                category=h["category"],
                severity=h["severity"],
                matched_keywords=h["matched_keywords"],
                snippet=h["snippet"],
                description=h["description"],
                suggestion=h["suggestion"],
                sort_order=h["sort_order"],
                risk_source=h.get("risk_source", "rule"),
            )
            for h in hits
        ]
        await self._repo.replace_risks(contract.id, user_id, risks)
        return risks

    async def _update_contract_counts(self, contract: Contract, hits: list[dict[str, Any]]) -> None:
        """更新合同风险统计。"""
        contract.risk_count = len(hits)
        contract.high_count = sum(1 for h in hits if h["severity"] == "high")
        contract.medium_count = sum(1 for h in hits if h["severity"] == "medium")
        contract.low_count = sum(1 for h in hits if h["severity"] == "low")
        contract.updated_by = contract.created_by
        await self._session.flush()

    async def _get_or_404(self, user_id: int, contract_id: int) -> Contract:
        contract = await self._repo.get_contract(user_id, contract_id)
        if contract is None:
            raise BizException(10001, "Not Found", http_status=404)
        return contract

    async def _audit(
        self, action: str, method: str, path: str, body: dict[str, Any], operator_id: int,
        request_meta: dict[str, Any],
    ) -> None:
        await self._audit_repo.add_log(
            user_id=operator_id, module="contract", action=action, method=method, path=path,
            request_body=body, response_code=0, ip=request_meta.get("ip"),
            user_agent=request_meta.get("user_agent"), duration_ms=request_meta.get("duration_ms"),
            trace_id=request_meta.get("trace_id"),
        )